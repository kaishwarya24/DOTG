"""
Utility functions for document processing
"""

import os
from typing import Dict, List, Any

def parse_pdf(file_path: str) -> Dict[str, Any]:
    """Extract text from PDF files"""
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "content": text,
                "pages": len(reader.pages),
                "type": "pdf"
            }
    except ImportError:
        return {"error": "PyPDF2 not installed. Install: pip install PyPDF2"}
    except Exception as e:
        return {"error": f"Error parsing PDF: {str(e)}"}

def parse_docx(file_path: str) -> Dict[str, Any]:
    """Extract text from DOCX files"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        return {
            "content": text,
            "paragraphs": len(doc.paragraphs),
            "type": "docx"
        }
    except ImportError:
        return {"error": "python-docx not installed. Install: pip install python-docx"}
    except Exception as e:
        return {"error": f"Error parsing DOCX: {str(e)}"}

def parse_text(file_path: str) -> Dict[str, Any]:
    """Extract text from plain text files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        return {
            "content": text,
            "lines": len(text.split('\n')),
            "type": "text"
        }
    except Exception as e:
        return {"error": f"Error parsing text file: {str(e)}"}

def parse_document(file_path: str) -> Dict[str, Any]:
    """Auto-detect and parse document"""
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext == '.docx':
        return parse_docx(file_path)
    elif ext in ['.txt', '.md']:
        return parse_text(file_path)
    else:
        return {"error": f"Unsupported file type: {ext}"}

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Split text into overlapping chunks"""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        chunks.append(chunk)
    
    return chunks

def extract_keywords(text: str, top_n: int = 10) -> List[str]:
    """Extract top keywords from text (simple frequency-based)"""
    # Remove common stopwords
    stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                 'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'be', 
                 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 
                 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that'}
    
    words = text.lower().split()
    word_freq = {}
    
    for word in words:
        word = word.strip('.,!?;:()[]{}\"\'')
        if word and word not in stopwords and len(word) > 3:
            word_freq[word] = word_freq.get(word, 0) + 1
    
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    return [word for word, freq in sorted_words[:top_n]]

def create_metadata(text: str, chunk_id: int) -> Dict[str, Any]:
    """Create metadata for a text chunk"""
    return {
        "chunk_id": chunk_id,
        "word_count": len(text.split()),
        "char_count": len(text),
        "keywords": extract_keywords(text, 5),
        "difficulty": "medium" 
    }




def extract_correct_answer(question_text):
    """Extract the correct answer from generated question text"""
    import re
    
    # Look for patterns like "Correct answer: A" or "Answer: A"
    patterns = [
        r'correct\s+answer[:\s]*([ABCD])',
        r'answer[:\s]*([ABCD])',
        r'correct[:\s]*([ABCD])',
        r'\*\*([ABCD])\*\*',  # Bold answer
        r'([ABCD])\s*is\s*correct',
        r'option\s*([ABCD])\s*is\s*correct',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, question_text, re.IGNORECASE)
        if match:
            return match.group(1).upper()
    
    # Debug: print question text if no answer found
    print(f"\nDEBUG: No correct answer found in question text:")
    print(question_text[:500])
    
    # Return None instead of defaulting to A
    return None

def extract_reasoning(question_text):
    """Extract reasoning/explanation from generated question text"""
    import re
    
    # Look for explanation after "Explanation:" 
    explanation_match = re.search(r'explanation[:\s]*(.+?)(?=\n\s*question|$)', question_text, re.IGNORECASE | re.DOTALL)
    if explanation_match:
        explanation = explanation_match.group(1).strip()
        # Clean up the explanation
        explanation = re.sub(r'\n+', ' ', explanation)  # Replace multiple newlines with space
        explanation = re.sub(r'\s+', ' ', explanation)   # Replace multiple spaces with single space
        return explanation
    
    # Fallback patterns
    patterns = [
        r'reasoning[:\s]*(.+?)(?=\n\s*question|$)',
        r'why[:\s]*(.+?)(?=\n\s*question|$)',
        r'because[:\s]*(.+?)(?=\n\s*question|$)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, question_text, re.IGNORECASE | re.DOTALL)
        if match:
            reasoning = match.group(1).strip()
            reasoning = re.sub(r'\n+', ' ', reasoning)
            reasoning = re.sub(r'\s+', ' ', reasoning)
            return reasoning
    
    return "No explanation provided for this question."

def clean_question_for_display(question_text):
    """Remove correct answer and explanation information from question before showing to user"""
    import re
    
    # Split by "Correct answer:" to separate question from answer/explanation
    parts = re.split(r'correct\s+answer[:\s]*', question_text, flags=re.IGNORECASE)
    
    if len(parts) > 1:
        # Take only the part before "Correct answer:"
        cleaned_text = parts[0].strip()
    else:
        # Fallback: remove patterns if no clear split
        cleaned_text = question_text
        patterns_to_remove = [
            r'answer[:\s]*[ABCD].*$',
            r'the\s+correct\s+answer\s+is[:\s]*.*$',
            r'\([ABCD]\)\s*correct.*$',
            r'\*\*[ABCD]\*\*.*$',
            r'[ABCD]\)\s*\*.*$',
            r'explanation[:\s]*.*$',
            r'reasoning[:\s]*.*$',
        ]
        
        for pattern in patterns_to_remove:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE | re.MULTILINE | re.DOTALL)
    
    # Remove any trailing whitespace and empty lines
    cleaned_text = re.sub(r'\n\s*\n', '\n', cleaned_text)
    cleaned_text = cleaned_text.strip()
    
    # Ensure all 4 options are present
    option_pattern = r'[ABCD]\)\s*.+'
    options_found = re.findall(option_pattern, cleaned_text)
    
    if len(options_found) < 4:
        # If less than 4 options found, return original text with warning
        return f"[WARNING: Only {len(options_found)} options found]\n{cleaned_text}"
    
    return cleaned_text